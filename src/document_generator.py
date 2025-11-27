"""
文档生成器模块
负责生成Docx和Markdown格式的文档
"""

import logging
from pathlib import Path
from typing import List, Any
from dataclasses import dataclass

from .logging_config import get_logger

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    get_logger(__name__).warning("python-docx库未安装，Docx生成功能将受限")

logger = get_logger(__name__)


class DocumentGenerator:
    """文档生成器类"""
    
    def __init__(self):
        self.logger = get_logger(__name__)
    
    def generate_docx(self, processed_slides: List[Any], output_path: Path):
        """
        生成Docx文档
        
        Args:
            processed_slides: 处理后的幻灯片数据
            output_path: 输出文件路径
        """
        self.logger.info(f"生成Docx文档: {output_path}")
        
        try:
            # 创建文档
            doc = Document()
            
            # 设置文档属性
            doc.core_properties.title = "PPT转换文档"
            doc.core_properties.author = "AI PPT Converter"
            doc.core_properties.subject = "PPT转换"
            
            # 添加标题
            title = doc.add_heading("PPT转换文档", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元信息
            self._add_metadata(doc, processed_slides)
            
            # 添加目录
            doc.add_heading("目录", 1)
            for i, slide in enumerate(processed_slides, 1):
                doc.add_paragraph(f"{i}. {slide.title}", style="List Number")
            
            doc.add_page_break()
            
            # 添加幻灯片内容
            for slide in processed_slides:
                self._add_slide_to_docx(doc, slide)
            
            # 保存文档
            doc.save(str(output_path))
            self.logger.info(f"Docx文档生成成功: {output_path}")
            
        except Exception as e:
            self.logger.error(f"生成Docx文档失败: {e}")
            raise
    
    def generate_markdown(self, processed_slides: List[Any], output_path: Path):
        """
        生成Markdown文档
        
        Args:
            processed_slides: 处理后的幻灯片数据
            output_path: 输出文件路径
        """
        self.logger.info(f"生成Markdown文档: {output_path}")
        
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                # 添加标题
                f.write("# PPT转换文档\n\n")
                
                # 添加元信息
                self._add_metadata_markdown(f, processed_slides)
                
                # 添加目录
                f.write("## 目录\n\n")
                for i, slide in enumerate(processed_slides, 1):
                    f.write(f"{i}. [{slide.title}](#slide-{i})\n")
                f.write("\n")
                
                # 添加幻灯片内容
                for slide in processed_slides:
                    self._add_slide_to_markdown(f, slide)
            
            self.logger.info(f"Markdown文档生成成功: {output_path}")
            
        except Exception as e:
            self.logger.error(f"生成Markdown文档失败: {e}")
            raise
    
    def _add_metadata(self, doc: Any, processed_slides: List[Any]):
        """添加元信息到Docx文档"""
        doc.add_heading("文档信息", 1)
        
        metadata = doc.add_paragraph()
        metadata.add_run("生成时间: ").bold = True
        metadata.add_run(f"{self._get_current_time()}\n")
        
        metadata.add_run("幻灯片数量: ").bold = True
        metadata.add_run(f"{len(processed_slides)}\n")
        
        metadata.add_run("AI服务: ").bold = True
        if processed_slides:
            metadata.add_run(f"{processed_slides[0].metadata.get('ai_service', '未知')}\n")
        
        metadata.add_run("模型: ").bold = True
        if processed_slides:
            metadata.add_run(f"{processed_slides[0].metadata.get('model', '未知')}\n")
        
        doc.add_paragraph()
    
    def _add_metadata_markdown(self, f, processed_slides: List[Any]):
        """添加元信息到Markdown文档"""
        f.write("## 文档信息\n\n")
        f.write(f"- **生成时间**: {self._get_current_time()}\n")
        f.write(f"- **幻灯片数量**: {len(processed_slides)}\n")
        
        if processed_slides:
            f.write(f"- **AI服务**: {processed_slides[0].metadata.get('ai_service', '未知')}\n")
            f.write(f"- **模型**: {processed_slides[0].metadata.get('model', '未知')}\n")
        
        f.write("\n")
    
    def _add_slide_to_docx(self, doc: Any, slide: Any):
        """添加幻灯片内容到Docx文档"""
        # 添加幻灯片标题
        doc.add_heading(f"{slide.slide_index}. {slide.title}", 1)
        
        # 添加摘要
        if slide.summary:
            doc.add_heading("摘要", 2)
            doc.add_paragraph(slide.summary)
        
        # 添加主要内容
        if slide.content:
            doc.add_heading("主要内容", 2)
            self._add_formatted_text(doc, slide.content)
        
        # 添加关键点
        if slide.key_points:
            doc.add_heading("关键点", 2)
            for point in slide.key_points:
                doc.add_paragraph(point, style="List Bullet")
        
        # 添加标签
        if slide.tags:
            doc.add_heading("标签", 2)
            tags_text = ", ".join(slide.tags)
            doc.add_paragraph(tags_text)
        
        # 添加原始文本（如果需要）
        if slide.metadata.get("original_text") and slide.metadata["original_text"] != slide.content:
            doc.add_heading("原始文本", 2)
            self._add_formatted_text(doc, slide.metadata["original_text"])
        
        # 添加分页符
        doc.add_page_break()
    
    def _add_slide_to_markdown(self, f, slide: Any):
        """添加幻灯片内容到Markdown文档"""
        # 添加幻灯片标题
        f.write(f"## <a name=\"slide-{slide.slide_index}\"></a>{slide.slide_index}. {slide.title}\n\n")
        
        # 添加摘要
        if slide.summary:
            f.write("### 摘要\n\n")
            f.write(f"{slide.summary}\n\n")
        
        # 添加主要内容
        if slide.content:
            f.write("### 主要内容\n\n")
            f.write(f"{slide.content}\n\n")
        
        # 添加关键点
        if slide.key_points:
            f.write("### 关键点\n\n")
            for point in slide.key_points:
                f.write(f"- {point}\n")
            f.write("\n")
        
        # 添加标签
        if slide.tags:
            f.write("### 标签\n\n")
            f.write(f"{', '.join(slide.tags)}\n\n")
        
        # 添加原始文本（如果需要）
        if slide.metadata.get("original_text") and slide.metadata["original_text"] != slide.content:
            f.write("### 原始文本\n\n")
            f.write(f"{slide.metadata['original_text']}\n\n")
        
        f.write("---\n\n")
    
    def _add_formatted_text(self, doc: Any, text: str):
        """添加格式化文本到Docx文档"""
        # 按段落分割
        paragraphs = text.split("\n")
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
                
            para = doc.add_paragraph()
            
            # 检查是否为标题
            if para_text.startswith("#"):
                level = min(para_text.count("#"), 6)
                heading_text = para_text.lstrip("#").strip()
                doc.add_heading(heading_text, level)
            # 检查是否为项目符号
            elif para_text.startswith("- ") or para_text.startswith("* "):
                bullet_text = para_text[2:].strip()
                doc.add_paragraph(bullet_text, style="List Bullet")
            # 检查是否为编号列表
            elif para_text[0].isdigit() and para_text[1:3] in [". ", ") "]:
                list_text = para_text[3:].strip()
                doc.add_paragraph(list_text, style="List Number")
            else:
                # 普通段落
                para.add_run(para_text)
    
    def _get_current_time(self) -> str:
        """获取当前时间"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


# 辅助函数
def create_table_from_data(doc: Any, data: List[List[str]]):
    """从数据创建表格"""
    if not data:
        return
    
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.cell(i, j).text = cell_text
