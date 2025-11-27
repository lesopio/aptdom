#!/usr/bin/env python3
"""
Docx生成器模块
负责将PPT内容转换为Word文档
"""

import logging
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass

from ppt_parser import SlideContent

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    logging.getLogger(__name__).warning("python-docx库未安装，Docx生成功能将受限")

logger = logging.getLogger(__name__)


class DocxGenerator:
    """Docx生成器类"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def generate(self, slides: List[SlideContent], output_path: Path) -> None:
        """
        生成Word文档
        
        Args:
            slides: 幻灯片内容列表
            output_path: 输出文件路径
        """
        self.logger.info(f"开始生成Docx文档: {output_path}")
        
        try:
            # 尝试使用python-docx库生成
            self._generate_with_docx(slides, output_path)
        except Exception as e:
            self.logger.warning(f"使用python-docx生成失败: {e}")
            # 回退到基本文本生成
            self._generate_basic_docx(slides, output_path)
    
    def _generate_with_docx(self, slides: List[SlideContent], output_path: Path) -> None:
        """使用python-docx库生成文档"""
        document = Document()
        
        # 设置默认样式
        style = document.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(12)
        
        # 添加标题
        title = document.add_heading('PPT转换文档', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 处理每页幻灯片
        for slide in slides:
            # 添加幻灯片标题
            if slide.title:
                document.add_heading(slide.title, level=1)
            
            # 添加文本内容
            if slide.text_content:
                document.add_paragraph(slide.text_content)
            
            # 添加项目符号
            for bullet in slide.bullet_points:
                document.add_paragraph(bullet, style='List Bullet')
            
            # 添加表格
            for table_data in slide.tables:
                self._add_table_to_docx(document, table_data)
            
            # 添加备注
            if slide.notes:
                document.add_heading('备注', level=2)
                document.add_paragraph(slide.notes)
            
            # 添加分页符（最后一页除外）
            if slide != slides[-1]:
                document.add_page_break()
        
        # 保存文档
        document.save(output_path)
        self.logger.info(f"成功生成Docx文档: {output_path}")
    
    def _add_table_to_docx(self, document, table_data: Dict[str, Any]) -> None:
        """添加表格到文档"""
        table = document.add_table(
            rows=len(table_data["data"]),
            cols=len(table_data["data"][0]) if table_data["data"] else 0
        )
        
        # 填充表格内容
        for i, row_data in enumerate(table_data["data"]):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text
        
        # 设置表格样式
        table.style = 'Table Grid'
    
    def _generate_basic_docx(self, slides: List[SlideContent], output_path: Path) -> None:
        """基本文本生成（当python-docx不可用时）"""
        self.logger.warning("使用基本文本生成模式，功能受限")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("PPT转换文档\n")
            f.write("=" * 50 + "\n\n")
            
            for slide in slides:
                if slide.title:
                    f.write(f"{slide.title}\n")
                    f.write("-" * len(slide.title) + "\n")
                
                if slide.text_content:
                    f.write(f"{slide.text_content}\n\n")
                
                for bullet in slide.bullet_points:
                    f.write(f"• {bullet}\n")
                
                if slide.notes:
                    f.write(f"\n备注: {slide.notes}\n")
                
                f.write("\n" + "=" * 50 + "\n\n")
        
        self.logger.info(f"生成基本文本文档: {output_path}")
